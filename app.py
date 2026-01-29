import os
import pandas as pd
import pytesseract
import pdfplumber
from flask import Flask, render_template, request, send_file, jsonify
from werkzeug.utils import secure_filename
from pdf2image import convert_from_path
from PIL import Image, ImageOps, ImageFilter
from docx import Document
from datetime import datetime

app = Flask(__name__)

# Config
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB
ALLOWED_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg', 'docx', 'txt'}

for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
    os.makedirs(folder, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def preprocess_image(image):
    """
    Improves OCR accuracy for handwriting by converting to grayscale
    and applying thresholding.
    """
    image = image.convert('L')  # Grayscale
    image = image.filter(ImageFilter.SHARPEN)
    return image

def extract_ocr(image, lang='eng+hin'):
    """Performs OCR with specific handwriting/language config"""
    # PSM 3: Fully automatic page segmentation, but no orientation detection.
    # OEM 3: Default, based on what is available.
    custom_config = r'--oem 3 --psm 3'
    return pytesseract.image_to_string(image, lang=lang, config=custom_config)

def process_pdf(path, lang):
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            content = page.extract_text()
            if content and len(content.strip()) > 20: # If text exists
                text += content + "\n"
            else:
                # If page is an image or handwriting, use OCR
                # We convert only the specific page to image to save memory
                images = convert_from_path(path, first_page=page.page_number, last_page=page.page_number)
                for img in images:
                    processed_img = preprocess_image(img)
                    text += extract_ocr(processed_img, lang) + "\n"
    return text

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['file']
    target_format = request.form.get('format', 'txt')
    lang_choice = request.form.get('lang', 'eng+hin') # Default to both

    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    filename = secure_filename(file.filename)
    input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(input_path)

    try:
        ext = filename.rsplit('.', 1)[1].lower()
        extracted_text = ""

        if ext == 'pdf':
            extracted_text = process_pdf(input_path, lang_choice)
        elif ext in ['jpg', 'jpeg', 'png']:
            img = Image.open(input_path)
            extracted_text = extract_ocr(preprocess_image(img), lang_choice)
        elif ext == 'docx':
            doc = Document(input_path)
            extracted_text = "\n".join([p.text for p in doc.paragraphs])
        else:
            with open(input_path, 'r', encoding='utf-8') as f:
                extracted_text = f.read()

        if not extracted_text.strip():
            return jsonify({'error': 'Could not detect any text. Ensure the image is clear.'}), 400

        # Save Logic
        base_name = filename.rsplit('.', 1)[0]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_filename = f"{base_name}_{timestamp}.{target_format}"
        out_path = os.path.join(app.config['OUTPUT_FOLDER'], out_filename)

        if target_format == 'txt':
            with open(out_path, 'w', encoding='utf-8') as f: f.write(extracted_text)
        elif target_format == 'md':
            with open(out_path, 'w', encoding='utf-8') as f: f.write(f"# OCR Result\n\n{extracted_text}")
        elif target_format == 'docx':
            d = Document()
            d.add_paragraph(extracted_text)
            d.save(out_path)
        elif target_format == 'csv':
            lines = [l for l in extracted_text.split('\n') if l.strip()]
            pd.DataFrame(lines, columns=['Content']).to_csv(out_path, index=False, encoding='utf-8-sig')

        return jsonify({
            'success': True,
            'preview': extracted_text[:1500],
            'download_url': f'/download/{out_filename}',
            'filename': out_filename
        })

    except Exception as e:
        print(f"Error: {e}")
        return jsonify({'error': str(e)}), 500
    finally:
        if os.path.exists(input_path):
            os.remove(input_path)

@app.route('/download/<filename>')
def download(filename):
    return send_file(os.path.join(app.config['OUTPUT_FOLDER'], filename), as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True, port=5000, threaded=True)
